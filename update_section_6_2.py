import json, os, re
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOCX = Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v5.0.docx")
LOG = Path(r"D:\GK\T\data\t-system\t_system_2026-09-02.log")
TMP = DOCX.with_name(DOCX.stem + ".tmp.docx")

rows = [
    ("订单/单据标识", "so、so1、sf_no、quotation_ref", "是", "部分", "quotation_ref 可映射为 Quotation Ref / PO No 并参与标准化；so、so1、sf_no 作为业务标识原样传递，不做记忆替换。"),
    ("流程与状态", "date、prior、stage、original", "是", "否", "日期、优先标志、阶段和原始单标志属于流程控制字段，按接口约定校验格式后原样传递。"),
    ("区域与条款", "term、location", "是", "是", "可做枚举映射、大小写和空格标准化；确认稳定映射后可配置记忆规则。"),
    ("人员姓名", "buyer、buyer_boss、ratifier、salesman、sales_person、applicant、presales", "是", "是", "可统一姓名写法或映射标准人员名称；人员身份和权限仍由 ePortal 校验。"),
    ("人员账号与联系方式", "user_name、buyer_mail、buyer_boss_mail、ratifier_mail、applicant_mail、applicant_id", "是", "否", "账号、邮件和人员 ID 用于识别或审计，应原样传递，不允许通过记忆规则改写。"),
    ("客户主数据", "customer_id、customer_name", "是", "部分", "customer_name 可标准化并作为记忆锚定键；customer_id 是主数据标识，只传输不转换。"),
    ("收货与最终用户", "customer_address、user_contact、delivery_date", "是", "部分", "地址和最终用户文本可做清洗、拆分及记忆替换；delivery_date 只校验日期格式，不做记忆替换。"),
    ("税务与付款", "tax_structure、customer_payment_term", "是", "是", "映射为 Tax Structure、Customer Payment Term；可做百分比、账期格式标准化和记忆替换。"),
    ("销售与业务分类", "sales_bundling、es_salesman_code", "是", "是", "可映射 ePortal 标准枚举或销售代码；未建立映射时保留智眸原值并记录未转换原因。"),
    ("汇率", "exchange_rate", "是", "否", "按数值格式传输，由 ePortal 按业务规则校验；不通过记忆规则修改。"),
    ("订单金额与 GP", "total_gp、product_gp、service_gp、total_amount、total_revenue、product_amount、service_amount、product_revenue、service_revenue", "是", "否", "智眸可传入作为参考值；金额、收入和 GP 属于计算/校验字段，ePortal 为最终计算真源并以只读方式展示。"),
    ("产品行容器", "products", "是", "是", "接收产品数组并转换为 ePortal 标准产品行结构；每个产品字段按下列规则处理。"),
    ("产品标识", "products[].product_id、products[].PN、products[].node_id", "是", "部分", "PN 可做料号格式标准化或映射；product_id、node_id 作为系统标识原样传输。"),
    ("产品描述与备注", "products[].description、products[].notes、products[].remarks", "是", "是", "可清洗空格、符号和文本格式；description 可参与记忆替换，备注类字段转换须保留原始值。"),
    ("数量、单价与币种", "products[].qty、products[].unit_price、products[].price、products[].currency", "是", "部分", "qty、unit_price 做数值格式转换，price、currency 做币种/枚举映射；不得用记忆规则改变真实数量或价格。"),
    ("产品供应与物流", "products[].supplier、products[].warehouse、products[].dropship、products[].inventory_type、products[].biz_category", "是", "是", "可映射标准供应商、仓库、直送标志、库存类型和业务分类枚举。"),
    ("产品税率", "products[].tax_pyable", "是", "是", "按百分比/小数统一规则转换，并映射到产品行税率字段。"),
    ("产品成本与 GP", "products[].GP、products[].GP_percent、products[].unit_cost、products[].total_cost、products[].total_price", "是", "否", "可传入供比对，但属于计算/只读字段；由 ePortal 根据数量、价格、成本和税率重新计算。"),
]

def set_cell_margin(cell, top=55, left=70, bottom=55, right=70):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side, value in (("top",top),("left",left),("bottom",bottom),("right",right)):
        n = tcMar.find(qn("w:"+side))
        if n is None: n=OxmlElement("w:"+side); tcMar.append(n)
        n.set(qn("w:w"), str(value)); n.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); n=trPr.find(qn("w:tblHeader"))
    if n is None: n=OxmlElement("w:tblHeader"); trPr.append(n)
    n.set(qn("w:val"),"true")

def cant_split(row):
    trPr=row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None: trPr.append(OxmlElement("w:cantSplit"))

def keep_next(p):
    pPr=p._p.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is None: pPr.append(OxmlElement("w:keepNext"))

def set_table_geometry(table, widths):
    total=sum(widths)
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn("w:tblW"))
    if tblW is None: tblW=OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"),str(total)); tblW.set(qn("w:type"),"dxa")
    tblInd=tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd=OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"),"120"); tblInd.set(qn("w:type"),"dxa")
    layout=tblPr.find(qn("w:tblLayout"))
    if layout is None: layout=OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"),"fixed")
    grid=table._tbl.tblGrid
    for node in list(grid): grid.remove(node)
    for width in widths:
        node=OxmlElement("w:gridCol"); node.set(qn("w:w"),str(width)); grid.append(node)
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None: tcW=OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"),str(width)); tcW.set(qn("w:type"),"dxa")

# Verify that the matrix covers every field observed in the supplied log.
log_text=LOG.read_text(encoding="utf-8")
payload=json.loads(re.search(r"data=(\{.*\})", log_text).group(1))
top=set(payload); product=set(payload["products"][0])
matrix=" ".join(r[1] for r in rows)
assert all(k in matrix for k in top), sorted(k for k in top if k not in matrix)
assert all(("products[]."+k) in matrix for k in product), sorted(k for k in product if ("products[]."+k) not in matrix)

doc=Document(DOCX)
heading=None; old=None
for i,p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("6.2 标准订单字段"):
        heading=p
        # The current section contains one prose paragraph before Heading 1 chapter 7.
        for q in doc.paragraphs[i+1:]:
            if q.style and q.style.name == "Heading 1": break
            if q.text.strip(): old=q; break
        break
if heading is None or old is None: raise RuntimeError("6.2 section not found")

old.text=("根据智眸联调日志，当前请求包含 43 个订单级字段和 21 个产品行字段。"
          "下表中的“可转换”指 T 系统可执行字段映射、格式标准化或记忆规则替换；"
          "标识、账号、日期以及金额/GP 等计算字段仅传输或校验，不做记忆改写。")
old.style=doc.styles["Normal"]
old.paragraph_format.space_after=Pt(4)
keep_next(old)

table=doc.add_table(rows=1, cols=5)
table.style="Table Grid"; table.autofit=False
headers=["字段类别","智眸字段（日志实际）","可传输","可转换","处理说明"]
widths=[1250,3000,800,850,3460]
for j,(txt,w) in enumerate(zip(headers,widths)):
    c=table.rows[0].cells[j]; c.text=txt; c.width=w
for item in rows:
    cells=table.add_row().cells
    for j,(txt,w) in enumerate(zip(item,widths)):
        cells[j].text=txt; cells[j].width=w

set_table_geometry(table,widths)

set_repeat_header(table.rows[0])
for ri,row in enumerate(table.rows):
    cant_split(row)
    for ci,cell in enumerate(row.cells):
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_before=Pt(1)
            p.paragraph_format.space_after=Pt(1)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if (ri==0 or ci in (0,2,3)) else WD_ALIGN_PARAGRAPH.LEFT
            if ri==0: keep_next(p)
            for run in p.runs:
                run.font.name="Microsoft YaHei"
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"),"微软雅黑")
                run.font.size=Pt(8)
                if ri==0: run.bold=True

# Move table directly after the explanatory paragraph.
old._p.addnext(table._tbl)

doc.save(TMP)
os.replace(TMP,DOCX)
print(DOCX)
