from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

src = Document(Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0.docx"))
out = Document(Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0_表格格式已对齐.docx"))

assert len(src.tables) == len(out.tables) == 15
assert [p.text for p in src.paragraphs] == [p.text for p in out.paragraphs]

def table_text(t):
    return [[c.text for c in r.cells] for r in t.rows]

for i, (a, b) in enumerate(zip(src.tables, out.tables)):
    assert table_text(a) == table_text(b), f"table {i+1} text changed"
    assert (len(a.rows), len(a.columns)) == (len(b.rows), len(b.columns))
    ga = [x.get(qn('w:w')) for x in a._tbl.tblGrid.gridCol_lst]
    gb = [x.get(qn('w:w')) for x in b._tbl.tblGrid.gridCol_lst]
    assert ga == gb, f"table {i+1} grid changed"
    assert b.rows[0]._tr.trPr.find(qn('w:tblHeader')) is not None
    for ri, row in enumerate(b.rows):
        assert row._tr.trPr.find(qn('w:cantSplit')) is not None
        for ci, cell in enumerate(row.cells):
            assert cell.vertical_alignment is not None, (i, ri, ci, 'vAlign')
            mar = cell._tc.tcPr.find(qn('w:tcMar'))
            assert mar is not None, (i, ri, ci, 'margin')
            expected = (('top','55'),('left','90'),('bottom','55'),('right','90')) if i < 3 else (('top','80'),('left','120'),('bottom','80'),('right','120'))
            got = tuple((s, mar.find(qn('w:'+s)).get(qn('w:w'))) for s,_ in expected)
            assert got == expected, (i, ri, ci, got)
            for p in cell.paragraphs:
                if ri == 0 or (i >= 3 and ci == 0):
                    assert int(p.alignment) == 1, (i, ri, ci, 'center')
                elif i >= 3:
                    assert int(p.alignment) == 0, (i, ri, ci, 'left')
print('PASS: content unchanged; 15 tables verified')
