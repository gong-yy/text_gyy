from collections import Counter
from docx import Document
from docx.oxml.ns import qn

DOC = r"D:\GK\T\data\t-system\T系统需求规格说明书v5.0.docx"
doc = Document(DOC)
print('tables', len(doc.tables), 'paragraphs', len(doc.paragraphs), 'sections', len(doc.sections))

def val(el, xpath, attr=None):
    xs = el.xpath(xpath)
    if not xs:
        return None
    return xs[0].get(attr) if attr else xs[0].text

for ti, t in enumerate(doc.tables):
    grid = [x.get(qn('w:w')) for x in t._tbl.tblGrid.gridCol_lst] if t._tbl.tblGrid is not None else []
    widths=[]; vas=[]; aligns=[]; margins=[]; pspaces=[]; fonts=[]; sizes=[]
    for row in t.rows:
        for c in row.cells:
            widths.append(val(c._tc, './w:tcPr/w:tcW', qn('w:w')))
            vas.append(val(c._tc, './w:tcPr/w:vAlign', qn('w:val')))
            tcMar=c._tc.tcPr.first_child_found_in('w:tcMar')
            if tcMar is not None:
                margins.append(tuple((side, (tcMar.find(qn('w:'+side)).get(qn('w:w')) if tcMar.find(qn('w:'+side)) is not None else None)) for side in ('top','left','bottom','right')))
            else: margins.append(None)
            for p in c.paragraphs:
                aligns.append(p.alignment)
                pf=p.paragraph_format
                pspaces.append((pf.space_before.pt if pf.space_before else None,pf.space_after.pt if pf.space_after else None,pf.line_spacing))
                for r in p.runs:
                    if r.text.strip():
                        fonts.append(r.font.name)
                        sizes.append(r.font.size.pt if r.font.size else None)
    txt=' | '.join(c.text.replace('\n',' / ')[:35] for c in t.rows[0].cells)
    print(f'T{ti+1}: {len(t.rows)}x{len(t.columns)} style={t.style.name if t.style else None} autofit={t.autofit} grid={grid} head={txt}')
    print('  widths', Counter(widths).most_common(8), 'vAlign', Counter(vas), 'align', Counter(aligns), 'margins', Counter(margins).most_common(4))
    print('  spacing', Counter(pspaces).most_common(5), 'fonts', Counter(fonts).most_common(5), 'sizes', Counter(sizes).most_common(5))
