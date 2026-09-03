from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

src = Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0.docx")
doc = Document(src)
assert len(doc.tables) == 15

headings = [p for p in doc.paragraphs if p.style and p.style.name in {"Heading 1", "Heading 2", "Heading 3"}]
assert headings
for p in headings:
    assert p._p.get_or_add_pPr().find(qn("w:keepNext")) is not None, p.text

fr01 = None
for table in doc.tables:
    assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
    for ri, row in enumerate(table.rows):
        assert row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
        for cell in row.cells:
            assert cell._tc.get_or_add_tcPr().find(qn("w:vAlign")) is not None
            assert cell._tc.get_or_add_tcPr().find(qn("w:tcMar")) is not None
            if ri == 0:
                for p in cell.paragraphs:
                    assert p._p.get_or_add_pPr().find(qn("w:keepNext")) is not None
        if row.cells[0].text.strip() == "FR-01":
            fr01 = row.cells[2].text.strip()

expected = "客户名为空时不应用记忆规则，所有字段保持智眸原值，并在响应或日志中明确标记未匹配原因。"
assert fr01 == expected, fr01
assert not src.with_name(src.stem + '.tmp.docx').exists()
print(f"PASS: {len(headings)} headings, 15 tables, FR-01 updated")
