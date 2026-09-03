import os
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0.docx")
TMP = SRC.with_name(SRC.stem + ".tmp.docx")

def ensure(parent, tag):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    return node

def set_cell_margins(cell, top, left, bottom, right):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tcMar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    node = ensure(row._tr.get_or_add_trPr(), "w:tblHeader")
    node.set(qn("w:val"), "true")

def prevent_row_split(row):
    ensure(row._tr.get_or_add_trPr(), "w:cantSplit")

def set_keep_next(paragraph, value=True):
    pPr = paragraph._p.get_or_add_pPr()
    node = pPr.find(qn("w:keepNext"))
    if value and node is None:
        pPr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        pPr.remove(node)

def replace_cell_text_preserve_format(cell, new_text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        p.add_run(new_text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

doc = Document(SRC)

# Keep every numbered heading with the content that follows so headings do not
# become isolated at the bottom of a page (including 4.2 T2 人工修改).
for p in doc.paragraphs:
    if p.style and p.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
        set_keep_next(p, True)

for ti, table in enumerate(doc.tables):
    table.autofit = False
    set_repeat_header(table.rows[0])
    for ri, row in enumerate(table.rows):
        prevent_row_split(row)
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ti < 3:
                set_cell_margins(cell, 55, 90, 55, 90)
            else:
                set_cell_margins(cell, 80, 120, 80, 120)
            for p in cell.paragraphs:
                if ti < 3:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                else:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                if ri == 0 or (ti >= 3 and ci == 0):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif ti >= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Keep each table header attached to at least the first data row.
                if ri == 0:
                    set_keep_next(p, True)

# 4.1 / FR-01: make the acceptance point match the stated behavior instead of
# introducing an unsupported mandatory HTTP 400 response.
target_table = None
for table in doc.tables:
    if len(table.columns) == 3 and any(row.cells[0].text.strip() == "FR-01" for row in table.rows):
        target_table = table
        break
if target_table is None:
    raise RuntimeError("FR-01 table not found")
for row in target_table.rows:
    if row.cells[0].text.strip() == "FR-01":
        replace_cell_text_preserve_format(
            row.cells[2],
            "客户名为空时不应用记忆规则，所有字段保持智眸原值，并在响应或日志中明确标记未匹配原因。",
        )
        break

doc.save(TMP)
os.replace(TMP, SRC)
print(SRC)
