import os
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOCX=Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v5.0.docx")
TMP=DOCX.with_name(DOCX.stem+".tmp.docx")

def set_margins(cell,top=80,left=120,bottom=80,right=120):
    tcPr=cell._tc.get_or_add_tcPr()
    tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side,value in (("top",top),("left",left),("bottom",bottom),("right",right)):
        node=tcMar.find(qn("w:"+side))
        if node is None:
            node=OxmlElement("w:"+side); tcMar.append(node)
        node.set(qn("w:w"),str(value)); node.set(qn("w:type"),"dxa")

doc=Document(DOCX)
target=None
header=["字段类别","智眸字段（日志实际）","可传输","可转换","处理说明"]
for table in doc.tables:
    if [c.text.strip() for c in table.rows[0].cells]==header:
        target=table; break
if target is None: raise RuntimeError("6.2 table not found")

target.style="Table Grid"; target.autofit=False
for ri,row in enumerate(target.rows):
    for ci,cell in enumerate(row.cells):
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_before=Pt(2)
            p.paragraph_format.space_after=Pt(2)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if (ri==0 or ci in (0,2,3)) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name="微软雅黑"
                rPr=run._element.get_or_add_rPr(); fonts=rPr.get_or_add_rFonts()
                fonts.set(qn("w:ascii"),"微软雅黑")
                fonts.set(qn("w:hAnsi"),"微软雅黑")
                fonts.set(qn("w:eastAsia"),"微软雅黑")
                run.font.size=Pt(8.5)
                run.bold=(ri==0)

doc.save(TMP)
os.replace(TMP,DOCX)
print(DOCX)
