from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0.docx")
OUT = Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v4.0_表格格式已对齐.docx")

def set_cell_margins(cell, top, left, bottom, right):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tcMar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    node = trPr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        trPr.append(node)
    node.set(qn("w:val"), "true")

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))

doc = Document(SRC)
for ti, table in enumerate(doc.tables):
    table.autofit = False
    set_repeat_header(table.rows[0])
    for ri, row in enumerate(table.rows):
        prevent_row_split(row)
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # The first three front-matter tables already use a tighter compact grid.
            if ti < 3:
                set_cell_margins(cell, 55, 90, 55, 90)
            else:
                set_cell_margins(cell, 80, 120, 80, 120)
            for p in cell.paragraphs:
                if ti < 3:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                else:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                # Header rows and compact identifier columns align centrally;
                # narrative columns retain left alignment for readability.
                if ri == 0 or (ti >= 3 and ci == 0):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif ti >= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.save(OUT)
print(OUT)
