import json, re, zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

docx=Path(r"D:\GK\T\data\t-system\T系统需求规格说明书v5.0.docx")
log=Path(r"D:\GK\T\data\t-system\t_system_2026-09-02.log")
doc=Document(docx)

target=None
for table in doc.tables:
    if table.rows and [c.text.strip() for c in table.rows[0].cells] == ["字段类别","智眸字段（日志实际）","可传输","可转换","处理说明"]:
        target=table; break
assert target is not None, "6.2 field matrix missing"
assert len(target.rows)==19 and len(target.columns)==5
assert target.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
assert [x.get(qn("w:w")) for x in target._tbl.tblGrid.gridCol_lst] == ['1250','3000','800','850','3460']
for row in target.rows:
    assert row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
    assert [c._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w")) for c in row.cells] == ['1250','3000','800','850','3460']

payload=json.loads(re.search(r"data=(\{.*\})",log.read_text(encoding='utf-8')).group(1))
matrix=' '.join(c.text for row in target.rows for c in row.cells)
missing_top=[k for k in payload if k not in matrix]
missing_product=[k for k in payload['products'][0] if ('products[].'+k) not in matrix]
assert not missing_top and not missing_product,(missing_top,missing_product)
assert '43 个订单级字段和 21 个产品行字段' in '\n'.join(p.text for p in doc.paragraphs)
with zipfile.ZipFile(docx) as z:
    bad=z.testzip(); assert bad is None,bad
print(f'PASS: 6.2 matrix 18 groups; all {len(payload)} top-level and {len(payload["products"][0])} product fields covered; DOCX package valid')
